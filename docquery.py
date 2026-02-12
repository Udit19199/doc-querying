import streamlit as st
from PyPDF2 import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
import os
import time
from sklearn.metrics.pairwise import cosine_similarity
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_ollama import ChatOllama, OllamaEmbeddings
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_core.documents import Document
import uuid
import shutil
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from docx import Document as DocxDocument
from google.api_core.exceptions import ResourceExhausted
import faiss

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None

try:
    from dataset_builder import build_eval_dataset
except ImportError:

    def build_eval_dataset():
        raise ImportError(
            "dataset_builder module not found. RAGAS evaluation will be disabled."
        )
        return None


from ragas.metrics import (
    answer_relevancy,
    faithfulness,
    context_precision,
    context_recall,
)
from langchain_text_splitters import CharacterTextSplitter
from ragas import evaluate
from datasets import Dataset

try:
    from eval_runner import run_eval
except ImportError:

    def run_eval():
        raise ImportError("eval_runner module not found.")
        return None


import inspect
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from io import BytesIO
import difflib
import re
from reportlab.platypus import PageBreak


class SafeHuggingFaceEmbeddings(HuggingFaceEmbeddings):
    def __init__(
        self,
        model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs=None,
        encode_kwargs={},
        **kwargs,
    ):
        if model_kwargs is None:
            model_kwargs = {}
        model_kwargs["device"] = "cpu"  # force CPU
        model = SentenceTransformer(model_name_or_path=model_name, **model_kwargs)
        super().__init__(
            model_name=model_name,
            model_kwargs=model_kwargs,
            encode_kwargs=encode_kwargs,
            **kwargs,
        )
        self.client = model


import requests

try:
    response = requests.get("http://localhost:11434/api/tags")
    if response.status_code != 200:
        st.warning("Ollama server not detected. Run `ollama serve` first.")
except Exception:
    st.warning("Could not connect to Ollama. Ensure it's running.")
# genai.configure(api_key=api_key)

EVAL_BUFFER = []


def log_for_eval(question: str, contexts, answer: str):
    context_texts = [
        doc.page_content if hasattr(doc, "page_content") else str(doc)
        for doc in contexts
    ]
    EVAL_BUFFER.append(
        {"question": question, "contexts": context_texts, "answer": answer}
    )


# @@
def run_eval():
    ds = build_eval_dataset()
    return evaluate(
        ds, metrics=[faithfulness, answer_relevancy, context_precision, context_recall]
    )


# ---------- Helper functions (kept your logic) ----------
def get_file_text(uploaded_files):
    text = ""
    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name.lower()
        if file_name.endswith(".pdf"):
            pdf_reader = PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                page_text = page.extract_text() or ""
                text += page_text
        elif file_name.endswith(".docx"):
            doc = DocxDocument(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif file_name.endswith(".txt"):
            text += uploaded_file.read().decode("utf-8")
        else:
            st.warning(f"Unsupported file type: {file_name}")
    return text


def get_text_chunks(text, chunk_strategy):
    if chunk_strategy == "Recursive":
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500, chunk_overlap=100
        )
        return text_splitter.split_text(text)
    elif chunk_strategy == "Fixed-Size":
        text_splitter = CharacterTextSplitter(
            separator="\n", chunk_size=400, chunk_overlap=100
        )
        return text_splitter.split_text(text)
    elif chunk_strategy == "Paragraph":
        paragraphs = text.split("\n\n")
        chunks = [p.strip() for p in paragraphs if p.strip()]
        if not chunks:
            st.warning(
                "No paragraphs found using double newlines. Falling back to Recursive splitter."
            )
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=500, chunk_overlap=100
            )
            return text_splitter.split_text(text)
        return chunks
    else:
        st.warning("Unknown chunking strategy selected. Defaulting to Recursive.")
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500, chunk_overlap=100
        )
        return text_splitter.split_text(text)


def build_and_save_faiss(text_chunks, embedding_model, index_type, folder_name):
    embeddings = get_embeddings(embedding_model)
    vectors = embeddings.embed_documents(text_chunks)
    vectors_np = np.array(vectors, dtype="float32")
    dim = vectors_np.shape[1]

    # ---- pick FAISS index type ----
    if index_type == "Flat":
        index = faiss.IndexFlatL2(dim)
    elif index_type == "IVF":
        nlist = min(100, len(vectors_np))
        quantizer = faiss.IndexFlatL2(dim)
        index = faiss.IndexIVFFlat(quantizer, dim, nlist)
        if not index.is_trained:
            index.train(vectors_np)
    elif index_type == "HNSW":
        index = faiss.IndexHNSWFlat(dim, 32)
        index.hnsw.efConstruction = 40
    else:
        index = faiss.IndexFlatL2(dim)

    index.add(vectors_np)

    # ---- make docstore ----
    docs = [Document(page_content=t) for t in text_chunks]
    index_to_doc_id = {i: str(uuid.uuid4()) for i in range(len(docs))}
    docstore = InMemoryDocstore({index_to_doc_id[i]: docs[i] for i in range(len(docs))})

    # ---- construct FAISS store (handle both old/new signatures) ----
    try:
        import inspect

        sig = inspect.signature(FAISS)
        params = set(sig.parameters.keys())
        if "embedding" in params:
            faiss_store = FAISS(
                embedding=embeddings,
                index=index,
                docstore=docstore,
                index_to_docstore_id=index_to_doc_id,
            )
        elif "embedding_function" in params:
            faiss_store = FAISS(
                embedding_function=embeddings,
                index=index,
                docstore=docstore,
                index_to_docstore_id=index_to_doc_id,
            )
        else:
            # super old versions may only support positional args
            faiss_store = FAISS(embeddings, index, docstore, index_to_doc_id)
    except Exception:
        # final fallback: try "embedding_function", then positional
        try:
            faiss_store = FAISS(
                embedding_function=embeddings,
                index=index,
                docstore=docstore,
                index_to_docstore_id=index_to_doc_id,
            )
        except TypeError:
            faiss_store = FAISS(embeddings, index, docstore, index_to_doc_id)

    # ---- save ----
    if os.path.exists(folder_name):
        shutil.rmtree(folder_name)
    faiss_store.save_local(folder_name)


def _unwrap_index(index):
    """
    Recursively unwrap common FAISS wrapper objects to return the underlying raw faiss.Index.
    Handles typical wrappers like IndexIDMap / IndexIDMap2 and composite indexes that expose
    sub-index lists (IndexShards / other composite containers).

    This is defensive: if nothing to unwrap, returns the original object.
    """
    try:
        # If wrapper exposes `.index` (IndexIDMap, IndexIDMap2, some wrappers)
        if (
            hasattr(index, "index")
            and getattr(index, "index") is not None
            and getattr(index, "index") is not index
        ):
            return _unwrap_index(getattr(index, "index"))

        # Common composite fields that may contain sub-index lists / single sub-index
        for attr_name in (
            "shards",
            "indexes",
            "sub_indexes",
            "components",
            "sub_index",
            "components_",
        ):
            if hasattr(index, attr_name):
                val = getattr(index, attr_name)
                if isinstance(val, (list, tuple)) and val:
                    return _unwrap_index(val[0])
                elif val is not None:
                    return _unwrap_index(val)

    except Exception:
        # If anything went wrong while trying to inspect the object, return it unchanged
        pass

    return index


def get_conversational_chain():
    style = st.session_state.get("prompt_style", "Detailed Answer")
    if style == "Detailed Answer":
        prompt_template = """You are a knowledgeable assistant. Based on the context below, provide a detailed and insightful answer to the question.

    - Prioritize clarity, structure, and reasoning.
    - Do not copy context directly—summarize or rephrase where possible.
    - Highlight key facts, relationships, or patterns that go beyond surface-level details.

    If context is insufficient, make a logical guess and clearly state the uncertainty.

Context:
{context}

Question:
{question}

Answer:"""
    elif style == "Concise Summary":
        prompt_template = """Provide a brief and clear summary answer to the question using the context.

Context:
{context}

Question:
{question}

Answer:"""
    elif style == "Bullet Points":
        prompt_template = """Answer the question using bullet points for clarity and structure. Base your answer on the context below.

Context:
{context}

Question:
{question}

Answer:"""
    elif style == "Explain Like I'm 5":
        prompt_template = """Explain the answer in a very simple and easy-to-understand way, like you're explaining it to a 5-year-old.

Context:
{context}

Question:
{question}

Answer:"""
    else:
        prompt_template = "Context:\n{context}\n\nQuestion:\n{question}\n\nAnswer:"

    temperature = st.session_state.get("temperature", 0.5)
    try:
        model = ChatOllama(
            model=st.session_state.get("generating_model", "llama3.2"),
            temperature=temperature,
        )

    except Exception as e:
        st.error(f"Error with Ollama model: {str(e)}")
        return None

    prompt = PromptTemplate(
        template=prompt_template, input_variables=["context", "question"]
    )
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    return chain


def get_similarity_score(answer, context_chunks, embedding_model_name=None):
    """
    Robust similarity in [0,1] between `answer` and `context_chunks`.
    Strategy:
      1) Quick exact substring / sentence match -> 1.0
      2) Try embeddings (prefer local models, fallback to cloud if needed)
      3) Fuzzy ratio (difflib) as final fallback
    Returns float 0..1
    """
    try:
        # normalize contexts to text list
        texts = []
        for c in context_chunks:
            try:
                # LangChain Document -> .page_content
                if hasattr(c, "page_content"):
                    texts.append(c.page_content or "")
                else:
                    texts.append(str(c))
            except Exception:
                texts.append(str(c))
        joined = " ".join(texts).strip()
        ans = (answer or "").strip()
        if not ans:
            return 0.0

        # 1) quick exact substring or sentence-level exact -> strong signal
        if any(ans in t for t in texts if t):
            return 1.0

        # check exact sentence-level presence
        sentences = []
        for t in texts:
            sentences += [s.strip() for s in re.split(r"(?<=[.!?])\s+", t) if s.strip()]
        for s in sentences:
            if ans in s:
                return 1.0

        # 2) embeddings-based similarity (prefer local models)
        emb_attempts = []

        # prefer HuggingFace local (fast, avoids cloud quota); but honor requested model if explicitly asked
        use_hf_first = True
        if embedding_model_name == "OpenAI":
            use_hf_first = False
        if use_hf_first:
            try:
                emb = SafeHuggingFaceEmbeddings(
                    model_name="sentence-transformers/all-MiniLM-L6-v2"
                )
                embs = emb.embed_documents([joined, ans])
                score = cosine_similarity([embs[0]], [embs[1]])[0][0]
                # normalize to 0..1 (if negative map to (score+1)/2)
                if score < 0:
                    score = (score + 1.0) / 2.0
                return round(max(0.0, min(1.0, float(score))), 4)
            except Exception as e:
                emb_attempts.append(("hf_failed", str(e)))

        # if requested OpenAI explicitly or HF failed, try Ollama embeddings
        if (
            embedding_model_name == "Ollama (Nomic)"
            or embedding_model_name == "OpenAI"
            or not use_hf_first
        ):
            try:
                emb = OllamaEmbeddings(model="nomic-embed-text")
                embs = emb.embed_documents([joined, ans])
                score = cosine_similarity([embs[0]], [embs[1]])[0][0]
                if score < 0:
                    score = (score + 1.0) / 2.0
                return round(max(0.0, min(1.0, float(score))), 4)
            except Exception as e:
                emb_attempts.append(("ollama_failed", str(e)))

        # 3) final fallback: difflib fuzzy match across sentences / contexts
        best = 0.0
        for t in texts:
            try:
                r = difflib.SequenceMatcher(None, ans, t).ratio()
                if r > best:
                    best = r
            except Exception:
                continue
        return round(max(0.0, min(1.0, float(best))), 4)

    except Exception as e:
        st.warning(f"Similarity scoring failed (final fallback): {e}")
        return 0.0


def extract_reference_from_contexts(answer, contexts):
    """
    Pick the single best sentence from `contexts` that most closely matches `answer`.
    Used as a cheap 'reference' for RAGAS when no ground-truth reference is available.
    """
    texts = []
    for c in contexts:
        if hasattr(c, "page_content"):
            texts.append(c.page_content or "")
        else:
            texts.append(str(c or ""))

    # split into sentences
    sents = []
    for t in texts:
        sents.extend([s.strip() for s in re.split(r"(?<=[.!?])\s+", t) if s.strip()])

    if not sents:
        # fallback: join contexts (trim to reasonable length)
        return (" ".join(texts)).strip()[:400]

    # pick the sentence with max fuzzy overlap to the answer
    best = max(
        sents, key=lambda s: difflib.SequenceMatcher(None, s, answer or "").ratio()
    )
    return best


def save_report_as_pdf(filename, eval_dict, metrics, folder_path="."):
    """
    Save query report as a PDF file.
    """
    file_path = os.path.join(folder_path, filename)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("<b>Query Evaluation Report</b>", styles["Title"]))
    story.append(Spacer(1, 12))

    story.append(
        Paragraph(f"<b>Question:</b> {eval_dict['question']}", styles["Normal"])
    )
    story.append(Spacer(1, 12))

    story.append(Paragraph("<b>Answer:</b>", styles["Heading3"]))
    story.append(Paragraph(eval_dict["answer"], styles["Normal"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("<b>Retrieved Contexts:</b>", styles["Heading3"]))
    for i, ctx in enumerate(eval_dict["contexts"]):
        story.append(Paragraph(f"Context {i + 1}: {ctx}", styles["Normal"]))
        story.append(Spacer(1, 6))

    story.append(Spacer(1, 12))
    story.append(Paragraph("<b>Evaluation Metrics:</b>", styles["Heading3"]))

    table_data = [["Metric", "Score"]]
    for k, v in metrics.items():
        table_data.append([k, f"{v:.2f}" if isinstance(v, (int, float)) else str(v)])

    table = Table(table_data, colWidths=[150, 150])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]
        )
    )
    story.append(table)

    doc = SimpleDocTemplate(file_path, pagesize=(595.27, 841.89))  # A4
    doc.build(story)
    return file_path


def compare_across_models(user_question, index_type, chunk_strategy):
    st.subheader("Comparing Responses Across Models:")
    raw_text = st.session_state.get("raw_text", "")
    text_chunks = get_text_chunks(raw_text, chunk_strategy)
    models_to_compare = {}
    try:
        models_to_compare["Ollama (Nomic)"] = OllamaEmbeddings(model="nomic-embed-text")
    except Exception as e:
        st.warning(f"Ollama Embeddings failed: {str(e)}")
    try:
        models_to_compare["MiniLM (HuggingFace)"] = SafeHuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
    except Exception as e:
        st.error(f"HuggingFace Embeddings failed: {str(e)}")
    if os.getenv("OPENAI_API_KEY"):
        try:
            models_to_compare["OpenAI"] = OpenAIEmbeddings()
        except Exception as e:
            st.warning(f"OpenAI Embeddings failed: {str(e)}")

    if not models_to_compare:
        st.warning("No embedding backends available for comparison.")
        return

    cols = st.columns(len(models_to_compare))
    for idx, (name, embedding) in enumerate(models_to_compare.items()):
        try:
            vectors = embedding.embed_documents(text_chunks)
            vectors_np = np.array(vectors, dtype="float32")
            dim = vectors_np.shape[1]
            if index_type == "Flat":
                index = faiss.IndexFlatL2(dim)
                index.add(vectors_np)
            elif index_type == "IVF":
                nlist = min(100, len(vectors_np))
                quantizer = faiss.IndexFlatL2(dim)
                index = faiss.IndexIVFFlat(quantizer, dim, nlist)
                index.train(vectors_np)
                index.add(vectors_np)
            elif index_type == "HNSW":
                index = faiss.IndexHNSWFlat(dim, 32)
                index.hnsw.efConstruction = 40
                index.add(vectors_np)
            else:
                index = faiss.IndexFlatL2(dim)
                index.add(vectors_np)

            db = FAISS.from_texts(text_chunks, embedding=embedding)
            tk = st.session_state.get("top_k", 5)
            docs = db.similarity_search(user_question, k=tk, fetch_k=max(20, tk * 5))
            chain = get_conversational_chain()
            if not chain:
                cols[idx].error(f"Chain failed for {name}")
                continue
            response = chain.invoke(
                {"input_documents": docs, "question": user_question},
                return_only_outputs=True,
            )
            cols[idx].markdown(f"**{name}**")
            cols[idx].write(response["output_text"])
        except Exception as e:
            cols[idx].error(f"{name} failed: {str(e)}")


def rerank_chunks_llm(question, chunks, top_k=None):
    if top_k is None:
        top_k = st.session_state.get("top_k", 5)

    model = ChatOllama(
        model=st.session_state.get("generating_model", "llama3.2"),
        temperature=0.2,
    )

    reranked = []
    for chunk in chunks:
        prompt = f"""
You are a smart AI assistant. Rate how relevant the following text chunk is to the question.

Chunk:
\"\"\"
{chunk.page_content}
\"\"\"

Question: {question}

Respond with a single number from 0 (not relevant) to 10 (highly relevant).
"""
        try:
            raw = model.invoke(prompt)
            text = getattr(raw, "content", raw)  # be defensive
            text = str(text).strip()
            score = float(text.split()[0])
            reranked.append((chunk, score))
        except Exception as e:
            print(f"Scoring failed: {e}")
            reranked.append((chunk, 0.0))

    reranked.sort(key=lambda x: x[1], reverse=True)
    top_chunks = [chunk for chunk, _ in reranked[:top_k]]
    return top_chunks


##
def compare_between_sets(user_question=None):

    cols = st.columns(2)
    eval_data = []
    chain = get_conversational_chain()
    for i, (label, set_key, faiss_folder) in enumerate(
        [
            ("Set A", "setA", "faiss_setA"),
            ("Set B", "setB", "faiss_setB"),
        ]
    ):
        try:
            if set_key not in st.session_state:
                raise ValueError(f"{set_key} not configured in session_state")
            emb_model_name = st.session_state[set_key]["embedding_model"]
            emb_model = get_embeddings(emb_model_name)
            db = FAISS.load_local(
                faiss_folder, emb_model, allow_dangerous_deserialization=True
            )
            top_k = st.session_state.get("top_k", 5)
            # ---------- DEBUG: show index size and top_k requested ----------
            try:
                index_size = getattr(db.index, "ntotal", None)
            except Exception:
                index_size = None
            # fallback to docstore size
            try:
                docstore_size = len(getattr(db, "docstore")._dict)
            except Exception:
                docstore_size = None

            base_index = _unwrap_index(db.index)

            if isinstance(base_index, faiss.IndexIVFFlat):
                nlist = getattr(base_index, "nlist", None) or 100
                probe = max(2, min(nlist, top_k * 2))
                base_index.nprobe = probe
            elif isinstance(base_index, faiss.IndexHNSWFlat):
                base_index.hnsw.efSearch = max(50, top_k * 4)

            query = user_question or ""
            docs = db.similarity_search(query, k=top_k, fetch_k=max(20, top_k * 5))

            if st.session_state.get("enable_rerank"):
                with cols[i]:
                    st.info(f"Re-ranking top {len(docs)} chunks using Ollama...")
                docs = rerank_chunks_llm(
                    user_question, docs, top_k=st.session_state.get("top_k", 5)
                )
            start_time = time.time()
            response = chain.invoke(
                {"input_documents": docs, "question": user_question},
                return_only_outputs=True,
            )
            end_time = time.time()
            latency = round(end_time - start_time, 2)

            answer_text = response["output_text"]

            if label == "Set A":
                st.session_state.eval_A = {
                    "question": user_question,
                    "contexts": [d.page_content for d in docs],  # strings for ragas
                    "answer": answer_text,
                }
            elif label == "Set B":
                st.session_state.eval_B = {
                    "question": user_question,
                    "contexts": [d.page_content for d in docs],
                    "answer": answer_text,
                }

            # Log for evaluation
            log_for_eval(user_question, docs, answer_text)
            similarity = get_similarity_score(answer_text, docs, emb_model_name)
            cols[i].markdown(f"**{label}**")
            cols[i].write(answer_text)

            eval_data.append(
                {
                    "Set": label,
                    "Embedding": emb_model_name,
                    "Index": st.session_state[set_key]["index_type"],
                    "Chunking": st.session_state[set_key]["chunk_strategy"],
                    "Answer Length": len(answer_text),
                    "Latency (s)": latency,
                    "Relevance Score": similarity,
                }
            )
        except Exception as e:
            cols[i].error(f"{label} failed: {str(e)}")

    # === RAGAS evaluation (unchanged logic, but we store results into session_state so we can export them) ===
    st.markdown("RAGAS Evaluation")

    colA, colB = st.columns(2)

    # Helper to format metric
    def fmt_metric(val):
        return "—" if pd.isna(val) else f"{float(val):.2f}"

    def my_llm_factory():
        return ChatOllama(model="llama3.2", temperature=0.0)

    def my_embedding_factory():
        return OllamaEmbeddings(model="nomic-embed-text")

    def run_ragas_on(eval_dict, label, container):
        ref_text = eval_dict.get("reference")
        if not ref_text:
            ref_text = extract_reference_from_contexts(
                eval_dict.get("answer", ""), eval_dict.get("contexts", [])
            )

        ds = Dataset.from_dict(
            {
                "question": [eval_dict.get("question", "")],
                "contexts": [eval_dict.get("contexts", [])],
                "answer": [eval_dict.get("answer", "")],
                "ground_truth": [ref_text],
            }
        )

        try:
            eval_llm = ChatOllama(
                model=st.session_state.get("generating_model", "llama3.2"),
                temperature=0.0,
            )
        except Exception as e:
            container.error(f"Could not instantiate evaluation LLM: {e}")
            return None

        try:
            eval_emb = OllamaEmbeddings(model="nomic-embed-text")
        except Exception as e:
            container.error(f"Could not instantiate evaluation embeddings: {e}")
            return None

        llm_for_ragas = None
        emb_for_ragas = None
        try:
            from ragas.llms import LangchainLLMWrapper

            llm_for_ragas = LangchainLLMWrapper(eval_llm)
        except Exception:
            try:
                from ragas.llms import LangchainLLM

                llm_for_ragas = LangchainLLM(eval_llm)
            except Exception:
                llm_for_ragas = eval_llm

        try:
            from ragas.embeddings import LangchainEmbeddingsWrapper

            emb_for_ragas = LangchainEmbeddingsWrapper(eval_emb)
        except Exception:
            emb_for_ragas = eval_emb

        sig = None
        try:
            sig = inspect.signature(evaluate)
            params = set(sig.parameters.keys())
        except Exception:
            params = set()

        metrics_list = [
            answer_relevancy,
            faithfulness,
            context_precision,
            context_recall,
        ]

        try:
            if "llm" in params or "embeddings" in params:
                results = evaluate(
                    ds,
                    metrics=metrics_list,
                    llm=llm_for_ragas,
                    embeddings=emb_for_ragas,
                )
                return results

            if "llm_or_chain_factory" in params:
                results = evaluate(
                    ds, llm_or_chain_factory=llm_for_ragas, metrics=metrics_list
                )
                return results

            if "llm_factory" in params or "embedding_factory" in params:
                results = evaluate(
                    ds,
                    metrics=metrics_list,
                    llm_factory=(lambda: llm_for_ragas),
                    embedding_factory=(lambda: emb_for_ragas),
                )
                return results

            try:
                results = evaluate(ds, metrics=metrics_list, llm=llm_for_ragas)
                return results
            except Exception:
                results = evaluate(ds, metrics=metrics_list)
                return results

        except ResourceExhausted as e:
            container.warning(
                f"Evaluation LLM quota issue for {label}: {str(e)}. Falling back to evaluate() without custom models."
            )
            try:
                results = evaluate(ds, metrics=metrics_list)
                return results
            except Exception as e2:
                container.error(f"Evaluation completely failed for {label}: {e2}")
                return None
        except TypeError as e:
            container.error(f"Evaluation failed due to unexpected signature: {e}")
            return None
        except Exception as e:
            container.error(f"Evaluation failed for {label}: {e}")
            return None

    # Run for Set A
    with colA:
        if "eval_A" in st.session_state:
            container = st.container()
            container.subheader("Set A")
            container.markdown(f"**Question:** {st.session_state.eval_A['question']}")
            with container.expander("Retrieved Chunks (A)", expanded=False):
                contexts = (
                    st.session_state.eval_A.get("contexts", [])
                    if "eval_A" in st.session_state
                    else []
                )
                if not contexts:
                    container.write("No chunks retrieved.")
                else:
                    sentinel = "— Select a chunk —"
                    options = [sentinel] + [
                        f"Chunk {i + 1}" for i in range(len(contexts))
                    ]
                    selected = container.selectbox(
                        "Preview a single chunk", options, index=0, key="select_chunk_A"
                    )
                    if selected != sentinel:
                        idx = options.index(selected) - 1
                        container.markdown(f"**Chunk {idx + 1}:**\n\n{contexts[idx]}")

                    if container.checkbox("Show all chunks", key="show_all_A"):
                        for j, c in enumerate(contexts, 1):
                            container.markdown(f"**Chunk {j}:**\n\n{c}\n\n---")

            container.markdown("**Answer:**")
            container.write(st.session_state.eval_A["answer"])

            res_a = run_ragas_on(st.session_state.eval_A, "Set A", container)
            # store ragas result for export later
            st.session_state["last_ragas_A"] = res_a

            if res_a is not None:
                df_a = res_a.to_pandas()

                def safe_get(df, key, default=None):
                    try:
                        val = df.at[0, key]
                        if pd.isna(val):
                            return default
                        return float(val)
                    except Exception:
                        return default

                # ---- fallback scores ----
                fallback_score = get_similarity_score(
                    st.session_state.eval_A["answer"],
                    [
                        Document(page_content=c)
                        for c in st.session_state.eval_A["contexts"]
                    ],
                    "Ollama (Nomic)",
                )

                answer_rel = safe_get(df_a, "answer_relevancy", fallback_score)
                faithful = safe_get(df_a, "faithfulness", fallback_score)
                ctx_prec = safe_get(df_a, "context_precision", fallback_score)
                ctx_rec = safe_get(df_a, "context_recall", fallback_score)

                container.metric(
                    "Answer Relevancy",
                    f"{answer_rel:.2f}",
                    help="How well the generated answer addresses the user’s question.",
                )
                container.metric(
                    "Faithfulness",
                    f"{faithful:.2f}",
                    help="Does the answer stay consistent with the retrieved context.",
                )
                container.metric(
                    "Context Precision",
                    f"{ctx_prec:.2f}",
                    help="Proportion of retrieved context that was actually relevant.",
                )
                container.metric(
                    "Context Recall",
                    f"{ctx_rec:.2f}",
                    help="How much of the relevant context was captured by the retrieval.",
                )
            else:
                container.warning(
                    "RAGAS returned no results. Showing similarity fallback."
                )

    # Run for Set B
    with colB:
        if "eval_B" in st.session_state:
            container = st.container()
            container.subheader("Set B")
            container.markdown(f"**Question:** {st.session_state.eval_B['question']}")
            with container.expander("Retrieved Chunks (B)", expanded=False):
                contexts = (
                    st.session_state.eval_B.get("contexts", [])
                    if "eval_B" in st.session_state
                    else []
                )
                if not contexts:
                    container.write("No chunks retrieved.")
                else:
                    sentinel = "— Select a chunk —"
                    options = [sentinel] + [
                        f"Chunk {i + 1}" for i in range(len(contexts))
                    ]
                    selected = container.selectbox(
                        "Preview a single chunk", options, index=0, key="select_chunk_B"
                    )
                    if selected != sentinel:
                        idx = options.index(selected) - 1
                        container.markdown(f"**Chunk {idx + 1}:**\n\n{contexts[idx]}")

                    if container.checkbox("Show all chunks", key="show_all_B"):
                        for j, c in enumerate(contexts, 1):
                            container.markdown(f"**Chunk {j}:**\n\n{c}\n\n---")

            container.markdown("**Answer:**")
            container.write(st.session_state.eval_B["answer"])

            res_b = run_ragas_on(st.session_state.eval_B, "Set B", container)
            # store ragas result for export later
            st.session_state["last_ragas_B"] = res_b

            #
            if res_b is not None:
                df_b = res_b.to_pandas()

                def safe_get(df, key, default=None):
                    try:
                        val = df.at[0, key]
                        if pd.isna(val):
                            return default
                        return float(val)
                    except Exception:
                        return default

                # ---- fallback scores ----
                fallback_score = get_similarity_score(
                    st.session_state.eval_B["answer"],
                    [
                        Document(page_content=c)
                        for c in st.session_state.eval_B["contexts"]
                    ],
                    "Ollama (Nomic)",
                )

                answer_rel = safe_get(df_b, "answer_relevancy", fallback_score)
                faithful = safe_get(df_b, "faithfulness", fallback_score)
                ctx_prec = safe_get(df_b, "context_precision", fallback_score)
                ctx_rec = safe_get(df_b, "context_recall", fallback_score)

                container.metric(
                    "Answer Relevancy",
                    f"{answer_rel:.2f}",
                    help="How well the generated answer addresses the user’s question.",
                )
                container.metric(
                    "Faithfulness",
                    f"{faithful:.2f}",
                    help="Does the answer stay consistent with the retrieved context.",
                )
                container.metric(
                    "Context Precision",
                    f"{ctx_prec:.2f}",
                    help="Proportion of retrieved context that was actually relevant.",
                )
                container.metric(
                    "Context Recall",
                    f"{ctx_rec:.2f}",
                    help="How much of the relevant context was captured by the retrieval.",
                )
            else:
                container.warning(
                    "RAGAS returned no results. Showing similarity fallback."
                )

    st.markdown("---")
    st.subheader("Export / Download Report")

    # Ensure default export choice
    if "export_choice" not in st.session_state:
        st.session_state["export_choice"] = "Both sets"

    # Dropdown (no form)
    export_choice = st.selectbox(
        "Which set do you want to export?",
        ["Set A", "Set B", "Both sets"],
        index=["Set A", "Set B", "Both sets"].index(st.session_state["export_choice"]),
        key="export_dropdown",
    )

    # Save the current selection
    st.session_state["export_choice"] = export_choice

    # --- Helper: Build metrics (same logic)
    def build_metrics_from_ragas(ragas_result, eval_dict, embedding_model_name):
        fallback = get_similarity_score(
            eval_dict.get("answer", ""),
            eval_dict.get("contexts", []),
            embedding_model_name,
        )
        default_metrics = {
            "Faithfulness": fallback,
            "Answer Relevancy": fallback,
            "Context Precision": fallback,
            "Context Recall": fallback,
        }

        if ragas_result is None:
            return default_metrics
        try:
            df = ragas_result.to_pandas()

            def safe_val(col):
                if col in df.columns and not pd.isna(df.at[0, col]):
                    return float(df.at[0, col])
                else:
                    return fallback

            return {
                "Faithfulness": safe_val("faithfulness"),
                "Answer Relevancy": safe_val("answer_relevancy"),
                "Context Precision": safe_val("context_precision"),
                "Context Recall": safe_val("context_recall"),
            }
        except Exception:
            return default_metrics

    # --- Helper: Generate PDF bytes
    def generate_pdf_bytes_from_payload(payload_dict):
        buf = BytesIO()
        styles = getSampleStyleSheet()
        story = []
        for label, (eval_d, metrics) in payload_dict.items():
            story.append(
                Paragraph(f"<b>{label} — Query Evaluation Report</b>", styles["Title"])
            )
            story.append(Spacer(1, 8))
            story.append(
                Paragraph(
                    f"<b>Question:</b> {eval_d.get('question', '')}", styles["Normal"]
                )
            )
            story.append(Spacer(1, 6))
            story.append(Paragraph("<b>Answer:</b>", styles["Heading3"]))
            story.append(Paragraph(eval_d.get("answer", ""), styles["Normal"]))
            story.append(Spacer(1, 8))
            story.append(Paragraph("<b>Retrieved Contexts:</b>", styles["Heading3"]))
            for i, ctx in enumerate(eval_d.get("contexts", [])):
                story.append(Paragraph(f"Context {i + 1}:", styles["Normal"]))
                story.append(Paragraph(ctx, styles["Normal"]))
                story.append(Spacer(1, 6))
            story.append(Spacer(1, 8))
            story.append(Paragraph("<b>Evaluation Metrics:</b>", styles["Heading3"]))
            table_data = [["Metric", "Score"]]
            for k, v in metrics.items():
                table_data.append(
                    [k, f"{v:.2f}" if isinstance(v, (int, float)) else str(v)]
                )
            table = Table(table_data, colWidths=[200, 200])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 12))
            story.append(PageBreak())

        doc = SimpleDocTemplate(buf, pagesize=(595.27, 841.89))
        doc.build(story)
        buf.seek(0)
        return buf.getvalue()

    # --- Generate payload directly from session (no recomputation)
    payload = {}

    if export_choice == "Set A":
        eval_A = st.session_state.get("eval_A")
        ragas_A = st.session_state.get("last_ragas_A")
        if eval_A:
            metrics_A = build_metrics_from_ragas(
                ragas_A, eval_A, st.session_state.setA["embedding_model"]
            )
            payload = {"Set A": (eval_A, metrics_A)}
        else:
            st.warning("No evaluation data available for Set A.")

    elif export_choice == "Set B":
        eval_B = st.session_state.get("eval_B")
        ragas_B = st.session_state.get("last_ragas_B")
        if eval_B:
            metrics_B = build_metrics_from_ragas(
                ragas_B, eval_B, st.session_state.setB["embedding_model"]
            )
            payload = {"Set B": (eval_B, metrics_B)}
        else:
            st.warning("No evaluation data available for Set B.")

    else:  # Both sets
        eval_A = st.session_state.get("eval_A")
        eval_B = st.session_state.get("eval_B")
        ragas_A = st.session_state.get("last_ragas_A")
        ragas_B = st.session_state.get("last_ragas_B")
        if eval_A and eval_B:
            metrics_A = build_metrics_from_ragas(
                ragas_A, eval_A, st.session_state.setA["embedding_model"]
            )
            metrics_B = build_metrics_from_ragas(
                ragas_B, eval_B, st.session_state.setB["embedding_model"]
            )
            payload = {"Set A": (eval_A, metrics_A), "Set B": (eval_B, metrics_B)}
        else:
            st.warning("Both Set A and Set B must be evaluated first before exporting.")

    # --- If valid payload, show download button
    if payload:
        pdf_bytes = generate_pdf_bytes_from_payload(payload)
        st.download_button(
            f"Download {export_choice} Report",
            data=pdf_bytes,
            file_name=f"{export_choice.replace(' ', '_')}_report.pdf",
            mime="application/pdf",
        )


def get_embeddings(embedding_model):
    if embedding_model == "Ollama (Nomic)":
        return OllamaEmbeddings(model="nomic-embed-text")
    elif embedding_model == "MiniLM (HuggingFace)":
        return SafeHuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
    elif embedding_model == "OpenAI":
        return OpenAIEmbeddings()
    else:
        st.warning("Unknown embedding model selected. Defaulting to Ollama (Nomic).")
        return OllamaEmbeddings(model="nomic-embed-text")


def parameter_topbar():
    # defaults once
    if "setA" not in st.session_state:
        st.session_state.setA = {
            "embedding_model": "Ollama (Nomic)",
            "index_type": "Flat",
            "chunk_strategy": "Recursive",
        }
    if "setB" not in st.session_state:
        st.session_state.setB = {
            "embedding_model": "MiniLM (HuggingFace)",
            "index_type": "Flat",
            "chunk_strategy": "Recursive",
        }
    if "compare_ready" not in st.session_state:
        st.session_state.compare_ready = False
    if "enable_rerank" not in st.session_state:
        st.session_state.enable_rerank = False
    if "top_k" not in st.session_state:
        st.session_state.top_k = 5
    if "temperature" not in st.session_state:
        st.session_state.temperature = 0.5
    if "prompt_style" not in st.session_state:
        st.session_state.prompt_style = "Detailed Answer"
    # NEW: generating model + vector DB (display-only single-option lists)
    if "generating_model" not in st.session_state:
        st.session_state.generating_model = "gpt-oss:20b"
    if "vector_db" not in st.session_state:
        st.session_state.vector_db = "FAISS"

    st.markdown("### ")

    # ---------- TOP ROW: Set A | Set B (side-by-side) ----------
    colA, colB = st.columns([1, 1])

    # ---- SET A (compact: Index+Chunking side-by-side, Embedding below) ----
    with colA:
        st.caption("SET A")
        rowA1, rowA2 = st.columns([1, 1])
        with rowA1:
            st.session_state.setA["index_type"] = st.selectbox(
                "Index (A)", ["Flat", "IVF", "HNSW"], key="indexA_top"
            )
        with rowA2:
            st.session_state.setA["chunk_strategy"] = st.selectbox(
                "Chunking (A)",
                ["Recursive", "Fixed-Size", "Paragraph"],
                key="chunkA_top",
            )
        # embedding on its own (below compact row)
        st.session_state.setA["embedding_model"] = st.selectbox(
            "Embedding (A)",
            ["Ollama (Nomic)", "MiniLM (HuggingFace)", "OpenAI"],
            key="embedA_top",
        )

    # ---- SET B (compact: Index+Chunking side-by-side, Embedding below) ----
    with colB:
        st.caption("SET B")
        rowB1, rowB2 = st.columns([1, 1])
        with rowB1:
            st.session_state.setB["index_type"] = st.selectbox(
                "Index (B)", ["Flat", "IVF", "HNSW"], key="indexB_top"
            )
        with rowB2:
            st.session_state.setB["chunk_strategy"] = st.selectbox(
                "Chunking (B)",
                ["Recursive", "Fixed-Size", "Paragraph"],
                key="chunkB_top",
            )
        st.session_state.setB["embedding_model"] = st.selectbox(
            "Embedding (B)",
            ["Ollama (Nomic)", "MiniLM (HuggingFace)", "OpenAI"],
            key="embedB_top",
        )

    # ---------- COMMON PARAMETERS (full-width below top row) ----------
    st.markdown("---")
    st.subheader("COMMON PARAMETERS")

    # two-column layout for primary common controls
    c1, c2 = st.columns([1, 1])
    with c1:
        # put Top-K and Temp horizontally so height is compact
        s1, s2 = st.columns([1, 1])
        with s1:
            st.session_state.top_k = st.slider(
                "Top-K", 3, 20, st.session_state.top_k, 1, key="topk_top"
            )
        with s2:
            st.session_state.temperature = st.slider(
                "Temp", 0.0, 1.0, st.session_state.temperature, 0.1, key="temp_top"
            )
    with c2:
        st.session_state.prompt_style = st.selectbox(
            "Style",
            [
                "Detailed Answer",
                "Concise Summary",
                "Bullet Points",
                "Explain Like I'm 5",
            ],
            key="style_top",
        )
        st.session_state.enable_rerank = st.checkbox(
            "Enable re-ranking", value=st.session_state.enable_rerank, key="rerank_top"
        )

    # NEW: Generating model + Vector DB (single-option placeholders) in one row
    gcol, vcol = st.columns([1, 1])
    with gcol:
        st.session_state.generating_model = st.selectbox(
            "Generating model",
            ["gpt-oss:20b"],
            key="gen_model_top",
        )
        st.caption("Generative model used for answer generation (display only).")
    with vcol:
        st.session_state.vector_db = st.selectbox(
            "Vector DB",
            ["FAISS"],
            key="vector_db_top",
        )
        st.caption("Vector DB used for retrieval (display only).")

    # Confirm button (same behaviour as before)
    if st.button("Confirm Parameter Sets", use_container_width=True):
        st.session_state.compare_ready = True
        st.toast("Parameter sets locked. You can now upload & process.", icon="✅")


# ---------- UI: minimal sidebar for uploads ----------
def sidebar_upload_section():
    with st.sidebar:
        st.header("Upload & Process")
        uploaded_files = st.file_uploader(
            "Upload PDF, DOCX, or TXT",
            type=["pdf", "docx", "txt"],
            accept_multiple_files=True,
            key="uploader_main",
        )

        if st.button("Submit & Process", key="process_main"):
            if not st.session_state.compare_ready:
                st.error("First click **Confirm Parameter Sets** in the top bar.")
                return
            if not uploaded_files:
                st.error("Please upload at least one file.")
                return

            with st.spinner("Processing documents..."):
                raw_text = get_file_text(uploaded_files)
                st.session_state["raw_text"] = raw_text

                # --- Process Set A ---
                chunksA = get_text_chunks(
                    raw_text, st.session_state.setA["chunk_strategy"]
                )
                # st.write(f"DEBUG: chunksA length = {len(chunksA)}")

                build_and_save_faiss(
                    chunksA,
                    st.session_state.setA["embedding_model"],
                    st.session_state.setA["index_type"],
                    "faiss_setA",
                )

                # --- Process Set B ---
                chunksB = get_text_chunks(
                    raw_text, st.session_state.setB["chunk_strategy"]
                )
                # st.write(f"DEBUG: chunksB length = {len(chunksB)}")

                build_and_save_faiss(
                    chunksB,
                    st.session_state.setB["embedding_model"],
                    st.session_state.setB["index_type"],
                    "faiss_setB",
                )

                st.success("Both FAISS indexes saved.")


def main():
    st.set_page_config(page_title="Document QA Comparison", layout="wide")
    topbar_height = 120
    st.title("Document-based QA system")

    parameter_topbar()

    sidebar_upload_section()

    st.markdown('<div class="main-content">', unsafe_allow_html=True)

    user_question = st.text_input(
        "Ask a question from the uploaded files", key="user_question_input"
    )
    if user_question and "raw_text" in st.session_state:
        compare_between_sets(user_question)
    elif "raw_text" in st.session_state:
        st.info(
            "Upload complete — ask a question to run comparisons between Set A and Set B."
        )

    st.markdown("</div>", unsafe_allow_html=True)


if __name__ == "__main__":
    main()
